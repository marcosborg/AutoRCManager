from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Manual_de_Procedimentos_e_Funcoes_Roles_AutoRCManager.docx")
BLUE = "1F4D78"
MID_BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F6F9"
GRAY = "5F6B76"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, 9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Atualizar índice no Word (botão direito > Atualizar campo)"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])


def add_para(doc, text="", bold_lead=None, italic=False, align=None, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, 11, bold=True, color=BLUE)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, 11, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, 11, italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_font(r, 11)


def add_steps(doc, steps):
    numbering = doc.part.numbering_part.element
    base_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == str(base_num_id))
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(step)
        set_font(r, 11)


def add_callout(doc, label, text, tone="info"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE if tone == "info" else ("FFF4CE" if tone == "warn" else "FDEDED"))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_font(r, 10.5, bold=True, color=BLUE if tone == "info" else (GOLD if tone == "warn" else RED))
    r = p.add_run(text)
    set_font(r, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, 9.5, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_font(r, 9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def role_section(doc, title, mission, scope, routines, procedures, controls, escalation, kpis):
    # The first role follows a full-page matrix; a forced break there creates an
    # empty page in LibreOffice. Heading keep-with-next already moves it cleanly.
    if not title.startswith("3. "):
        doc.add_page_break()
    doc.add_heading(title, level=1)
    add_callout(doc, "Missão do perfil:", mission)
    doc.add_heading("Âmbito funcional", level=2)
    add_bullets(doc, scope)
    doc.add_heading("Rotina recomendada", level=2)
    add_table(doc, ["Momento", "Ação obrigatória", "Resultado esperado"], routines, [1500, 4800, 3060])
    doc.add_heading("Procedimentos detalhados", level=2)
    for name, objective, steps, evidence in procedures:
        doc.add_heading(name, level=3)
        add_para(doc, "Objetivo: " + objective, bold_lead="Objetivo:")
        add_steps(doc, steps)
        add_para(doc, "Registo/evidência: " + evidence, bold_lead="Registo/evidência:")
    doc.add_heading("Controlos e proibições", level=2)
    add_bullets(doc, controls)
    doc.add_heading("Escalonamento", level=2)
    add_table(doc, ["Situação", "Ação imediata", "Escalar para"], escalation, [3100, 3600, 2660])
    doc.add_heading("Indicadores de execução", level=2)
    add_bullets(doc, kpis)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 18, 10),
        ("Heading 2", 13, MID_BLUE, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.625)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AUTORC MANAGER  |  MANUAL DE PROCEDIMENTOS E FUNÇÕES"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        set_font(run, 8.5, bold=True, color=GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Capa editorial
    add_para(doc, "MANUAL INTERNO", align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    for _ in range(4):
        add_para(doc, "", after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Manual de Procedimentos\ne Funções")
    set_font(r, 30, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Perfis Adm, Aux. gestão, Gestão e Marketing")
    set_font(r, 15, color=MID_BLUE)
    add_callout(doc, "Finalidade:", "Definir responsabilidades, rotinas, controlos, evidências e critérios de escalonamento no AutoRCManager.")
    for _ in range(3):
        add_para(doc, "", after=8)
    add_para(doc, "Auto RC Manager", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Versão 1.0  |  4 de agosto de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Documento interno — validar sempre as permissões atuais do sistema", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    doc.add_page_break()
    doc.add_heading("Controlo do documento", level=1)
    add_table(doc, ["Campo", "Conteúdo"], [
        ("Documento", "Manual de Procedimentos e Funções por Role"),
        ("Aplicação", "AutoRCManager"),
        ("Perfis abrangidos", "Adm; Aux. gestão; Gestão; Marketing"),
        ("Versão / data", "1.0 / 04-08-2026"),
        ("Periodicidade de revisão", "Semestral e sempre que existam alterações de permissões ou processos"),
        ("Responsável pela aprovação", "Administração / Direção"),
        ("Classificação", "Uso interno"),
    ], [2200, 7160])
    doc.add_heading("Como utilizar este manual", level=2)
    add_bullets(doc, [
        "Consultar primeiro as regras comuns e depois o capítulo correspondente ao perfil atribuído.",
        "Executar cada operação com informação completa, comprovativo anexado quando aplicável e histórico rastreável.",
        "Não contornar bloqueios do sistema nem utilizar credenciais de outro utilizador.",
        "Em caso de conflito entre este manual e a permissão efetiva no sistema, prevalece o bloqueio do sistema e o caso deve ser escalado ao Adm.",
    ])
    doc.add_heading("Índice", level=2)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')

    doc.add_page_break()
    doc.add_heading("1. Princípios gerais de operação", level=1)
    add_para(doc, "Este manual traduz o modelo de acessos do AutoRCManager em regras operacionais. Uma permissão técnica permite uma ação, mas não dispensa validação documental, segregação de funções, autorização hierárquica ou bom senso profissional.")
    doc.add_heading("1.1 Regras transversais", level=2)
    add_bullets(doc, [
        "Autenticidade: registar apenas factos confirmados; nunca inventar datas, valores, estados, fornecedores ou documentos.",
        "Completude: preencher campos obrigatórios e relevantes antes de avançar o processo.",
        "Rastreabilidade: usar observações, anexos e históricos para permitir reconstituição posterior da decisão.",
        "Confidencialidade: consultar e partilhar dados pessoais e financeiros apenas para fins de trabalho.",
        "Segregação: quem prepara uma operação sensível não deve aprová-la quando o processo exigir validação independente.",
        "Privilégio mínimo: não usar ações de edição ou eliminação quando a consulta é suficiente.",
        "Correção controlada: preferir retificação documentada; eliminar apenas duplicados ou registos manifestamente inválidos e com autorização.",
    ])
    doc.add_heading("1.2 Fluxo padrão de qualquer registo", level=2)
    add_steps(doc, [
        "Pesquisar primeiro para evitar duplicados, usando matrícula, VIN, nome, NIF, telefone, e-mail ou referência.",
        "Confirmar a identidade do registo e o contexto operacional antes de editar.",
        "Recolher os dados e documentos de origem; validar legibilidade, data, entidade e coerência de valores.",
        "Introduzir ou atualizar a informação no módulo correto, sem abreviaturas ambíguas.",
        "Rever o resumo do registo e confirmar valores, datas, estado, proprietário/responsável e anexos.",
        "Guardar e verificar a mensagem de sucesso ou o histórico criado.",
        "Comunicar a conclusão ao interveniente seguinte quando exista dependência ou aprovação.",
    ])
    doc.add_heading("1.3 Operações críticas", level=2)
    add_table(doc, ["Operação", "Risco principal", "Controlo mínimo"], [
        ("Eliminar registos", "Perda de histórico e referências", "Confirmar duplicado/erro, dependências e autorização"),
        ("Alterar valores financeiros", "Pagamento ou margem incorretos", "Conferir documento de origem e segundo controlo"),
        ("Aprovar lotes/pagamentos", "Compromisso financeiro indevido", "Validar total, itens, fornecedor, documentos e estado"),
        ("Alterar roles/permissões", "Acesso excessivo ou bloqueio operacional", "Pedido formal, aprovação e teste com pré-visualização"),
        ("Converter retoma", "Criação incorreta de viatura/stock", "Confirmar negócio, proprietário, avaliação e documentos"),
        ("Fechar/reabrir trabalhos", "Estado operacional falso", "Confirmar execução física, custos e justificação"),
    ], [2400, 3100, 3860])
    add_callout(doc, "Regra de ouro:", "Se a operação produzir efeito financeiro, legal, de acesso ou de eliminação, parar e obter validação quando faltar um comprovativo, uma aprovação ou informação essencial.", tone="warn")

    doc.add_heading("2. Matriz resumida de responsabilidades", level=1)
    add_para(doc, "A matriz seguinte é operacional e deve ser lida em conjunto com as permissões atuais. C = conduz; V = valida/aprova; A = apoia/prepara; Q = consulta quando necessário; — = não integra a função normal.")
    add_table(doc, ["Processo", "Adm", "Gestão", "Aux. gestão", "Marketing"], [
        ("Utilizadores, roles e permissões", "C/V", "C, mediante autorização", "—", "—"),
        ("Configurações e tabelas-base", "V", "C", "A", "A limitada"),
        ("Clientes", "V", "C", "C", "C, no âmbito comercial"),
        ("Viaturas e documentação", "V", "C", "C", "A/C conteúdo comercial"),
        ("Lotes e pagamentos de aquisição", "V", "C/V", "A/C", "—"),
        ("Consignações", "V", "C", "C", "Q"),
        ("Retomas", "V", "C/V", "C sem conversão", "C sem conversão"),
        ("Oficina, peças e serviços externos", "V", "C", "A/C", "Q"),
        ("Caixa e informação financeira", "V", "C", "A apenas onde autorizado", "—"),
        ("Leads, desempenho e IA", "V", "Q/V", "Q quando solicitado", "C/A"),
        ("Auditoria e controlo", "C/V", "C", "A", "Q"),
    ], [3200, 1100, 1500, 1860, 1700])

    admin_procs = [
        ("A. Gestão de utilizadores e acessos", "Garantir que cada pessoa possui apenas os acessos necessários à função.", [
            "Confirmar pedido, identidade, função, responsável hierárquico e data de início/fim.",
            "Pesquisar o utilizador; criar apenas se não existir e usar e-mail individual.",
            "Atribuir o role aprovado; evitar permissões avulsas sem justificação documentada.",
            "Usar a pré-visualização de role para confirmar o menu e os limites antes de comunicar o acesso.",
            "Registar a aprovação e instruir o utilizador a alterar a palavra-passe no primeiro acesso.",
            "Na saída ou mudança de função, retirar imediatamente o role anterior e rever sessões/acessos.",
        ], "Pedido aprovado, utilizador/role atualizado e registo de auditoria."),
        ("B. Revisão de permissões e auditoria", "Detetar acessos excessivos, alterações sensíveis e anomalias.", [
            "Rever semestralmente utilizadores ativos, roles atribuídos e contas sem uso conhecido.",
            "Comparar permissões dos roles com as responsabilidades deste manual.",
            "Consultar logs de auditoria em eliminações, alterações financeiras, estados e acessos.",
            "Investigar ações fora de padrão; preservar evidências e ouvir o responsável.",
            "Corrigir o acesso, documentar a decisão e comunicar alterações relevantes.",
        ], "Lista de revisão, evidências analisadas e decisão de manutenção/correção."),
        ("C. Aprovação de operações sensíveis", "Autorizar apenas operações completas, coerentes e justificadas.", [
            "Abrir a operação e rever o objeto, responsável, documentos e histórico.",
            "Conferir totais, datas, entidade, método de pagamento e enquadramento do negócio.",
            "Confirmar que não existem duplicados, alertas pendentes ou dependências por resolver.",
            "Aprovar apenas quando todos os critérios estão cumpridos; caso contrário devolver com motivo objetivo.",
            "Depois da aprovação, confirmar o novo estado e a respetiva rastreabilidade.",
        ], "Estado de aprovação, comentário/justificação e documentos associados."),
        ("D. Manutenção funcional e incidentes", "Restabelecer a operação sem perda de dados nem alterações improvisadas.", [
            "Reproduzir o erro e recolher utilizador, data/hora, módulo, registo e mensagem apresentada.",
            "Confirmar se é falta de permissão, dado inválido, estado incompatível ou falha técnica.",
            "Aplicar apenas correções funcionais autorizadas; não alterar diretamente dados sem procedimento controlado.",
            "Escalar falhas técnicas com passos de reprodução e impacto.",
            "Testar a resolução com o perfil afetado e encerrar o incidente com evidência.",
        ], "Registo do incidente, causa, correção, teste e data de encerramento."),
    ]
    role_section(doc, "3. Role Adm", "Administrar o sistema, assegurar controlo interno e validar operações de maior risco, mantendo disponibilidade, segurança e rastreabilidade.", [
        "Administração de utilizadores, roles, permissões, configurações e auditoria.",
        "Acesso transversal a clientes, viaturas, vendas, aquisições, oficina, peças, pagamentos, caixa, IA e leads.",
        "Validação de operações sensíveis, exceções, eliminações e alterações estruturais.",
        "Acompanhamento de alertas de gestão, desempenho e integridade da informação.",
    ], [
        ("Início do dia", "Rever alertas, falhas, pendências de aprovação e operações financeiras críticas.", "Prioridades atribuídas e riscos identificados."),
        ("Durante o dia", "Responder a acessos, exceções, bloqueios e validações; monitorizar alterações sensíveis.", "Operação desbloqueada com controlo."),
        ("Fim do dia", "Confirmar pendências críticas, caixa/transferências quando aplicável e incidentes abertos.", "Handover claro e sem riscos esquecidos."),
        ("Semanal", "Rever logs, utilizadores, estados anómalos, aprovações e indicadores.", "Plano corretivo e responsáveis definidos."),
    ], admin_procs, [
        "Nunca partilhar credenciais nem atribuir superadmin por conveniência.",
        "Não apagar histórico para ocultar erros; corrigir e justificar.",
        "Não aprovar uma operação preparada pelo próprio quando exista alternativa de segundo validador.",
        "Não alterar configurações em produção sem avaliar impacto nos registos existentes.",
        "Usar a manutenção técnica apenas quando for o administrador real e houver necessidade comprovada.",
    ], [
        ("Suspeita de fraude ou acesso indevido", "Preservar logs, limitar acesso e não confrontar sem coordenação.", "Administração/Direção"),
        ("Falha técnica com impacto geral", "Registar evidência e suspender operações afetadas.", "Suporte técnico/Desenvolvimento"),
        ("Divergência financeira", "Bloquear aprovação e reconciliar documentos.", "Direção financeira/Contabilidade"),
        ("Pedido de acesso excessivo", "Recusar provisoriamente e pedir justificação/aprovação.", "Responsável hierárquico"),
    ], [
        "Tempo de resolução de pedidos de acesso e incidentes.",
        "Percentagem de revisões de acesso concluídas no prazo.",
        "Número de exceções, eliminações e reaberturas sem justificação.",
        "Pendências críticas com mais de um dia útil.",
    ])

    mgmt_procs = [
        ("A. Gestão de viaturas e circuito de compra", "Manter o stock e o custo de aquisição completos e coerentes.", [
            "Pesquisar por matrícula e VIN antes de criar.",
            "Confirmar origem, empresa compradora, fornecedor, valor, datas, localização e estado geral.",
            "Criar ou rever a ficha da viatura e anexar documentos disponíveis.",
            "Constituir o lote quando aplicável, conferir itens, despesas de registo/reboque e pagamentos.",
            "Rever o total do lote e submeter/aprovar apenas com comprovativos e correspondência entre itens.",
            "Atualizar estados e observações à medida que o veículo avança no circuito.",
        ], "Ficha, documentos, histórico de estados, lote e pagamentos associados."),
        ("B. Reconciliação de clientes e negócio", "Assegurar coerência entre cliente, viatura, venda, pagamentos e documentos.", [
            "Validar dados do cliente e corrigir duplicados antes de associar o negócio.",
            "Conferir preço, financiamento, retoma, pagamentos, encargos e saldo.",
            "Comparar o registo com contratos, faturas, comprovativos e informação bancária.",
            "Resolver diferenças justificadas; escalar divergências materiais.",
            "Confirmar o estado final e os documentos necessários à entrega/transferência.",
        ], "Reconciliação do cliente, comprovativos e observação de fecho."),
        ("C. Coordenação de oficina, peças e serviços", "Garantir que trabalhos, peças e custos avançam com prioridade e controlo.", [
            "Rever viaturas em oficina e trabalhos pendentes por estado e data prevista.",
            "Validar necessidades de peças/serviços e evitar encomendas duplicadas.",
            "Acompanhar encomenda, receção, pagamento e associação à reparação.",
            "Conferir execução, custos e documentos antes do fecho.",
            "Atualizar o estado da viatura e comunicar atrasos que afetem venda ou entrega.",
        ], "Reparação, encomenda/receção, pagamento, serviço externo e estado atualizados."),
        ("D. Gestão de retomas e consignações", "Controlar veículos de terceiros e entradas de stock sem perda de titularidade ou valor.", [
            "Confirmar proprietário, identificação do veículo, avaliação e condições acordadas.",
            "Registar a retoma/consignação e anexar a documentação de suporte.",
            "Verificar estado, destino, período efetivo e histórico de alterações.",
            "Na conversão de retoma, rever todos os dados antes de criar/associar a viatura.",
            "Encerrar ou corrigir apenas com justificação e preservação do histórico.",
        ], "Registo de retoma/consignação, documentos, auditoria e eventual viatura convertida."),
    ]
    role_section(doc, "4. Role Gestão", "Coordenar e controlar os processos administrativos, financeiros e operacionais, garantindo que viaturas, clientes, compras, vendas e oficina evoluem com dados fiáveis.", [
        "Gestão ampla de clientes, viaturas, aquisições, vendas, lotes, retomas e estados.",
        "Coordenação de oficina, peças, pagamentos, receções e serviços externos.",
        "Gestão de tabelas-base, utilizadores/roles quando formalmente autorizado e auditoria.",
        "Acesso a caixa e transferência de caixa, respeitando segregação e comprovativos.",
    ], [
        ("Início do dia", "Rever dashboard, estados, entregas, oficina, pagamentos e alertas vencidos.", "Plano diário priorizado."),
        ("Durante o dia", "Validar dados, coordenar equipas e desbloquear pendências documentais.", "Processos avançam sem lacunas."),
        ("Fim do dia", "Reconciliar operações concluídas e registar pendências/responsáveis.", "Informação pronta para controlo."),
        ("Semanal", "Analisar stock parado, custos, oficina, consignações e reconciliações.", "Ações corretivas definidas."),
    ], mgmt_procs, [
        "Não criar ou alterar utilizadores/roles sem pedido e autorização formal.",
        "Não eliminar viaturas, vendas, clientes ou documentos para corrigir uma incongruência normal.",
        "Não aprovar lotes/pagamentos sem conferir documentos e totalizadores.",
        "Não transferir caixa sem origem, destino, motivo e confirmação do recetor.",
        "Não fechar uma reparação ou processo sem execução e documentação confirmadas.",
    ], [
        ("Margem, pagamento ou saldo divergente", "Suspender fecho e reconciliar linha a linha.", "Adm/Direção financeira"),
        ("Documento legal em falta", "Bloquear entrega/transferência afetada.", "Adm/Responsável do negócio"),
        ("Atraso de oficina compromete entrega", "Repriorizar e informar impacto/prazo.", "Adm/Direção comercial"),
        ("Erro estrutural ou permissão em falta", "Registar caso e não contornar o sistema.", "Adm/Suporte técnico"),
    ], [
        "Tempo médio por etapa desde aquisição até disponibilização/venda.",
        "Percentagem de processos com documentos completos.",
        "Stock e trabalhos parados acima do prazo definido.",
        "Diferenças de reconciliação e operações devolvidas por erro.",
    ])

    aux_procs = [
        ("A. Criação e atualização de fichas", "Registar corretamente clientes e viaturas, sem duplicação.", [
            "Pesquisar por identificadores fortes antes de criar.",
            "Conferir os dados com documento de origem e preencher contactos, proveniência e identificação.",
            "Na viatura, validar matrícula, VIN, marca/modelo, origem, empresa compradora, estado e localização.",
            "Anexar documentos nas áreas corretas e usar observações objetivas.",
            "Rever a ficha guardada e assinalar à Gestão qualquer campo crítico em falta.",
        ], "Ficha completa, anexos e indicação de pendências."),
        ("B. Preparação de lotes e pagamentos", "Preparar operações de aquisição para validação, com total correto e suporte documental.", [
            "Confirmar o lote/fornecedor e associar apenas as viaturas corretas.",
            "Registar valores por item, despesas aplicáveis e método/origem do pagamento.",
            "Anexar fatura, adjudicação ou comprovativo correspondente.",
            "Comparar soma dos itens, encargos e total do pagamento.",
            "Submeter para aprovação e não alterar depois da validação sem reabertura autorizada.",
        ], "Lote preparado, total conferido, comprovativos e estado de aprovação."),
        ("C. Encomendas, receções e pagamentos de peças", "Manter o ciclo da peça rastreável da necessidade ao pagamento.", [
            "Pesquisar encomendas existentes para a mesma viatura/reparação e peça.",
            "Criar a encomenda com fornecedor, referência, quantidade, preço e viatura/reparação.",
            "Atualizar previsão e estado; registar atrasos nas observações.",
            "Na receção, conferir quantidade, estado e documento; registar diferenças.",
            "Preparar o pagamento com fatura/comprovativo e comunicar anomalias à Gestão.",
        ], "Encomenda, receção e pagamento associados à reparação e ao fornecedor."),
        ("D. Consignações e retomas", "Registar corretamente entradas, destinos e condições sem executar decisões reservadas.", [
            "Validar intervenientes, viatura, titularidade e condições acordadas.",
            "Criar/editar a consignação ou retoma e anexar a evidência disponível.",
            "Atualizar destino, período e estado sempre que haja movimento confirmado.",
            "Rever o histórico antes de corrigir ou eliminar.",
            "Na retoma, preparar toda a informação para a Gestão; não converter em stock se o perfil não possuir essa autorização.",
        ], "Registo completo e pendência encaminhada para validação/conversão."),
    ]
    role_section(doc, "5. Role Aux. gestão", "Executar e preparar tarefas administrativas e operacionais com rigor documental, deixando os processos completos para validação pela Gestão ou pelo Adm.", [
        "Criação e manutenção de clientes, viaturas e documentação autorizada.",
        "Preparação de aquisições, lotes, pagamentos, consignações e retomas.",
        "Operação de encomendas, receções e pagamentos de peças e serviços externos.",
        "Manutenção de fornecedores, marcas, transportadoras, estados e dados-base no âmbito definido.",
        "Acesso a dados financeiros sensíveis apenas para execução das tarefas autorizadas.",
    ], [
        ("Início do dia", "Consultar tarefas, documentos recebidos, encomendas e pendências devolvidas.", "Fila de trabalho ordenada."),
        ("Durante o dia", "Registar documentos, atualizar estados e pedir validações atempadamente.", "Processos completos e rastreáveis."),
        ("Fim do dia", "Rever guardados, anexos, pagamentos preparados e pendências.", "Nada crítico fica sem responsável."),
        ("Semanal", "Pesquisar duplicados, registos incompletos e encomendas atrasadas.", "Lista de correções entregue à Gestão."),
    ], aux_procs, [
        "Não usar a possibilidade técnica de eliminar sem confirmar autorização da Gestão/Adm.",
        "Não aprovar a própria preparação quando for necessário segundo controlo.",
        "Não criar estados, fornecedores ou marcas duplicados; pesquisar antes.",
        "Não alterar valores para fazer coincidir totais sem documento justificativo.",
        "Não converter retomas quando o role não possui a permissão de conversão.",
        "Não divulgar informação financeira sensível fora do processo de trabalho.",
    ], [
        ("Documento ilegível ou incoerente", "Não concluir; pedir nova evidência.", "Gestão"),
        ("Diferença de quantidade/valor", "Registar a diferença e suspender pagamento.", "Gestão"),
        ("Possível duplicado", "Não eliminar nem fundir por iniciativa própria.", "Gestão/Adm"),
        ("Pedido fora das permissões", "Documentar e encaminhar, sem usar conta alheia.", "Adm"),
    ], [
        "Percentagem de registos devolvidos por falta/erro.",
        "Tempo de processamento de documentos e encomendas.",
        "Número de duplicados criados.",
        "Pendências sem atualização acima do prazo definido.",
    ])

    mkt_procs = [
        ("A. Qualificação e tratamento de leads", "Transformar contactos em oportunidades rastreáveis e acionáveis.", [
            "Confirmar origem, consentimento/contexto, nome, contacto, viatura/interesse e mensagem.",
            "Pesquisar cliente/lead existente para evitar duplicação.",
            "Registar ou atualizar informação e classificar prioridade/estado de forma objetiva.",
            "Efetuar o primeiro contacto dentro do SLA definido e registar resultado e próxima ação.",
            "Encaminhar oportunidade qualificada para o responsável comercial e acompanhar até aceitação.",
            "Fechar ou marcar sem interesse apenas com motivo registado.",
        ], "Lead atualizado, histórico de contacto, responsável e próxima ação."),
        ("B. Conteúdo de viaturas e qualidade comercial", "Garantir que a informação publicada é correta, completa e coerente com o stock.", [
            "Abrir a ficha da viatura e confirmar estado comercial antes de preparar conteúdo.",
            "Validar marca/modelo, versão, ano, quilometragem, combustível, equipamento, preço e condições.",
            "Selecionar imagens atuais, nítidas e coerentes com a viatura; não ocultar defeitos relevantes.",
            "Redigir descrição factual, sem promessas não autorizadas nem informação pessoal.",
            "Submeter alterações sensíveis de preço/condições à Gestão antes da publicação.",
            "Retirar ou corrigir conteúdo quando o estado da viatura mudar.",
        ], "Ficha revista, conteúdo aprovado e data/canal de publicação."),
        ("C. Gestão de dados de suporte comercial", "Manter clientes, viaturas, vendas/estados e retomas úteis ao trabalho comercial.", [
            "Pesquisar antes de criar cliente ou viatura.",
            "Atualizar apenas dados confirmados e relevantes para o seguimento comercial.",
            "Registar retomas recebidas com identificação, contacto, avaliação preliminar e documentos disponíveis.",
            "Encaminhar a retoma para validação; não converter quando o perfil não dispõe dessa autorização.",
            "Consultar estados de venda/reparação apenas para comunicar informação confirmada ao cliente.",
        ], "Dados comerciais atualizados e encaminhamento registado."),
        ("D. Campanhas e aprendizagem", "Melhorar a geração e conversão de procura com dados fiáveis.", [
            "Definir objetivo, público, mensagem, canal, período, orçamento autorizado e métrica antes da campanha.",
            "Usar referências de stock e condições validadas pela Gestão.",
            "Acompanhar leads por origem, tempo de resposta, qualificação e resultado.",
            "Analisar desempenho sem alterar dados operacionais para melhorar artificialmente indicadores.",
            "Documentar conclusões e propor ajustes de conteúdo, canal ou segmentação.",
        ], "Briefing, peças aprovadas, métricas por origem e relatório de aprendizagem."),
    ]
    role_section(doc, "6. Role Marketing", "Gerir informação e atividades comerciais de atração, qualificação e comunicação, mantendo dados de clientes e viaturas corretos e alinhados com a operação.", [
        "Tratamento de dados comerciais de clientes, viaturas, vendas/estados e retomas.",
        "Produção e atualização de informação para divulgação de stock e campanhas.",
        "Acompanhamento de contactos e oportunidades quando os módulos de lead estiverem disponíveis ao perfil.",
        "Consulta de reparações/estados quando necessária para comunicação comercial, sem interferir na execução da oficina.",
        "Acesso a configurações administrativas apenas se formalmente atribuído; não constitui função normal do Marketing.",
    ], [
        ("Início do dia", "Rever novos contactos, respostas pendentes, alterações de stock e campanhas ativas.", "Prioridades e SLA de contacto definidos."),
        ("Durante o dia", "Qualificar, contactar, atualizar dados e coordenar publicações com a Gestão.", "Oportunidades com próxima ação."),
        ("Fim do dia", "Fechar seguimentos, atualizar estados e sinalizar leads quentes/bloqueios.", "Pipeline comercial atualizado."),
        ("Semanal", "Analisar origem, tempo de resposta, qualidade e conversão; auditar anúncios ativos.", "Plano de otimização baseado em dados."),
    ], mkt_procs, [
        "Não publicar preço, disponibilidade, financiamento ou garantia sem validação atual.",
        "Não criar avaliações/depreciações/apreciações para influenciar artificialmente uma campanha.",
        "Não eliminar clientes, viaturas, vendas ou reparações como forma de limpar listas.",
        "Não aceder ou partilhar dados pessoais além do necessário para o contacto e seguimento.",
        "Não alterar estados operacionais da oficina sem confirmação do responsável.",
        "Roles e permissões eventualmente visíveis não devem ser geridos pelo Marketing sem autorização expressa.",
    ], [
        ("Reclamação ou pedido de remoção de dados", "Registar pedido e suspender contacto quando aplicável.", "Adm/Responsável de proteção de dados"),
        ("Preço/estado comercial divergente", "Suspender publicação e confirmar a ficha.", "Gestão"),
        ("Lead de elevado valor ou urgência", "Encaminhar com contexto e prazo.", "Responsável comercial/Gestão"),
        ("Possível informação enganosa", "Retirar provisoriamente e preservar evidência.", "Adm/Gestão"),
    ], [
        "Tempo até ao primeiro contacto e percentagem dentro do SLA.",
        "Taxa de qualificação e de conversão por origem.",
        "Percentagem de anúncios/fichas sem erro factual.",
        "Leads sem próxima ação ou sem atualização.",
    ])

    doc.add_page_break()
    doc.add_heading("7. Procedimentos comuns por processo", level=1)
    common = [
        ("7.1 Clientes", [
            "Pesquisar por NIF, telefone e e-mail; confirmar duplicados antes de criar.",
            "Registar nome legal, contactos, proveniência e observações objetivas.",
            "Associar viaturas, vendas, retomas e pagamentos ao cliente correto.",
            "Corrigir dados com base em documento ou confirmação do titular.",
            "Restringir a consulta e exportação ao mínimo necessário.",
        ]),
        ("7.2 Viaturas e estados", [
            "Usar matrícula e VIN como identificadores principais.",
            "Manter aquisição, empresa compradora, localização, estado geral/oficina e mês de IUC quando aplicável.",
            "Anexar DAV, registos, faturas, declarações e restantes documentos nas áreas próprias.",
            "Registar transferências de estado apenas após ocorrência confirmada.",
            "Antes de recuperar/eliminar, verificar vendas, reparações, pagamentos, lotes e histórico.",
        ]),
        ("7.3 Caixa e pagamentos", [
            "Identificar origem/destino, departamento, categoria, método, valor, data e motivo.",
            "Anexar comprovativo e confirmar se o pagamento está associado ao objeto correto.",
            "Não fracionar operações para evitar aprovação.",
            "Nas transferências, exigir confirmação de envio e receção.",
            "Reconciliar diferenças no próprio dia e escalar se não resolvidas.",
        ]),
        ("7.4 Oficina e peças", [
            "Associar todo o trabalho à viatura/reparação correta.",
            "Definir estado, responsável, datas e descrição clara do trabalho.",
            "Ligar peças e serviços externos ao trabalho que os originou.",
            "Conferir receções e custos antes de fechar ou dar saída.",
            "Não comunicar conclusão ao cliente enquanto o estado físico e o sistema não coincidirem.",
        ]),
        ("7.5 Calendário, alertas e handover", [
            "Criar tarefa com responsável, data/hora, contexto e ligação ao registo relevante.",
            "Atualizar ou concluir tarefas no momento da execução.",
            "Reatribuir formalmente quando houver ausência; não deixar tarefas sem dono.",
            "No handover, listar situação, última ação, bloqueio, próximo passo e prazo.",
        ]),
    ]
    for heading, bullets in common:
        doc.add_heading(heading, level=2)
        add_bullets(doc, bullets)

    doc.add_heading("8. Checklists de controlo", level=1)
    add_table(doc, ["Antes de criar/editar", "Antes de aprovar/fechar", "Antes de eliminar"], [
        ("☐ Pesquisei duplicados\n☐ Tenho fonte válida\n☐ Escolhi o módulo correto\n☐ Preenchi campos essenciais\n☐ Anexei comprovativos", "☐ Conferi valores/datas\n☐ Revisei documentos\n☐ Resolvi alertas\n☐ Confirmei autorização\n☐ Registei justificação", "☐ Confirmei o alvo\n☐ Verifiquei dependências\n☐ Tenho autorização\n☐ Existe alternativa de correção\n☐ A ação ficará auditável"),
    ], [3120, 3120, 3120])
    doc.add_heading("Checklist de fim do dia", level=2)
    add_bullets(doc, [
        "Todos os registos guardados e revistos.",
        "Pagamentos, lotes, caixa e operações sensíveis sem diferenças conhecidas.",
        "Pendências críticas atribuídas a uma pessoa e com prazo.",
        "Documentos recebidos anexados ao processo correto.",
        "Clientes/leads com próxima ação atualizada.",
        "Incidentes e bloqueios comunicados com evidência suficiente.",
    ])

    doc.add_heading("9. Matriz de escalonamento e tempos", level=1)
    add_table(doc, ["Prioridade", "Exemplo", "Resposta esperada", "Canal"], [
        ("P1 — Crítica", "Acesso indevido, fraude, indisponibilidade geral, pagamento errado iminente", "Imediata; suspender a operação afetada", "Telefone/mensagem direta + registo"),
        ("P2 — Alta", "Entrega bloqueada, documento legal em falta, erro financeiro material", "No próprio dia útil", "Gestão/Adm + tarefa"),
        ("P3 — Normal", "Correção de dados, permissão necessária, divergência não urgente", "Até 2 dias úteis", "Tarefa ou pedido documentado"),
        ("P4 — Melhoria", "Sugestão de processo, relatório ou alteração de interface", "Planeamento periódico", "Backlog/reunião de melhoria"),
    ], [1400, 4000, 2300, 1660])

    doc.add_heading("10. Anexo — referência de permissões", level=1)
    add_para(doc, "A referência abaixo resume a configuração observada em 04-08-2026. As permissões podem evoluir; a validação deve ser feita em Gestão de utilizadores > Roles/Permissões e, quando necessário, com a pré-visualização de role.")
    add_table(doc, ["Role", "Capacidades distintivas", "Limites/alertas operacionais"], [
        ("Adm", "Administração integral; IA/leads; oficina/pintura; caixa; aprovações; superadmin no role Admin.", "Aplicar segregação; manutenção apenas por administrador real; alto risco em eliminações e acessos."),
        ("Gestão", "Gestão ampla de negócio; lotes/aprovação; oficina; caixa/transferências; utilizadores e configurações.", "Sem superadmin; alterações de acessos exigem autorização; não consta acesso a leads/IA na configuração observada."),
        ("Aux. gestão", "Clientes/viaturas; documentos; lotes; consignações; peças/serviços; informação financeira sensível.", "Retomas sem conversão; sem caixa; sem gestão de utilizadores; eliminações devem ser excecionais e autorizadas."),
        ("Marketing", "Clientes, viaturas, vendas/estados, retomas e reparações; várias tabelas-base e gestão de utilizadores tecnicamente atribuídas.", "Permissões administrativas observadas excedem a função típica; recomenda-se revisão pelo Adm. Sem conversão de retomas e sem caixa."),
    ], [1500, 4300, 3560])
    add_callout(doc, "Recomendação de governação:", "Rever especialmente as permissões administrativas e de eliminação do role Marketing, bem como as permissões de gestão de utilizadores atribuídas a Gestão. A atribuição técnica não deve substituir autorização organizacional.", tone="warn")

    doc.add_heading("11. Termo de conhecimento", level=1)
    add_para(doc, "Declaro que li, compreendi e me comprometo a cumprir o presente Manual de Procedimentos e Funções, usando o AutoRCManager apenas no âmbito das responsabilidades e autorizações que me foram atribuídas.")
    add_para(doc, "Nome: ______________________________________________________________", after=12)
    add_para(doc, "Role: _______________________________   Data: ____ / ____ / ______", after=12)
    add_para(doc, "Assinatura: _________________________________________________________", after=12)
    add_para(doc, "Responsável/Aprovador: ______________________________________________", after=12)

    doc.core_properties.title = "Manual de Procedimentos e Funções — AutoRCManager"
    doc.core_properties.subject = "Roles Adm, Aux. gestão, Gestão e Marketing"
    doc.core_properties.author = "Auto RC Manager"
    doc.core_properties.comments = "Gerado a partir da configuração funcional observada no AutoRCManager em 04-08-2026."
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
